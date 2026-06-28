from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = Path("docs/API_REST_Pet_Grooming.docx")


GROUPS = [
    (
        "Auth API",
        "Autenticación, MFA, cierre de sesión y usuario actual.",
        [
            ("POST", "/api/v1/auth/login", "Login cliente"),
            ("POST", "/api/v1/auth/intranet-login", "Login intranet"),
            ("POST", "/api/v1/auth/mfa", "Validar MFA"),
            ("GET", "/api/v1/auth/me", "Usuario actual"),
            ("POST", "/api/v1/auth/logout", "Cerrar sesión"),
        ],
    ),
    (
        "Cliente API",
        "Perfil, mascotas y reservas propias del cliente.",
        [
            ("GET", "/api/v1/clientes/perfil", "Ver perfil"),
            ("PUT/PATCH", "/api/v1/clientes/perfil", "Actualizar perfil"),
            ("GET", "/api/v1/clientes/mascotas", "Listar mascotas"),
            ("POST", "/api/v1/clientes/mascotas", "Registrar mascota"),
            ("GET", "/api/v1/clientes/mascotas/{mascota}", "Ver mascota"),
            ("PUT/PATCH", "/api/v1/clientes/mascotas/{mascota}", "Actualizar mascota"),
            ("DELETE", "/api/v1/clientes/mascotas/{mascota}", "Eliminar mascota"),
            ("GET", "/api/v1/clientes/reservas", "Listar reservas del cliente"),
            ("POST", "/api/v1/clientes/reservas", "Crear reserva"),
            ("GET", "/api/v1/clientes/reservas/{reserva}", "Ver reserva"),
            ("PUT/PATCH", "/api/v1/clientes/reservas/{reserva}", "Actualizar reserva"),
            ("DELETE", "/api/v1/clientes/reservas/{reserva}", "Cancelar reserva"),
        ],
    ),
    (
        "Reservas API",
        "Gestión de reservas, disponibilidad, pagos y boletas.",
        [
            ("GET", "/api/v1/reservas", "Listar reservas"),
            ("POST", "/api/v1/reservas", "Crear reserva"),
            ("GET", "/api/v1/reservas/{reserva}", "Ver reserva"),
            ("PUT/PATCH", "/api/v1/reservas/{reserva}", "Actualizar reserva"),
            ("DELETE", "/api/v1/reservas/{reserva}", "Cancelar reserva"),
            ("POST", "/api/v1/reservas/horarios-disponibles", "Consultar horarios"),
            ("GET", "/api/v1/reservas/{reserva}/pagos", "Ver pagos"),
            ("POST", "/api/v1/reservas/{reserva}/pagos", "Registrar pago"),
            ("POST", "/api/v1/reservas/{reserva}/pago", "Registrar pago"),
            ("GET", "/api/v1/reservas/{reserva}/boleta", "Ver boleta"),
            ("GET", "/api/v1/reservas/{reserva}/boleta/descargar", "Descargar boleta"),
            ("GET", "/api/v1/pagos/{pago}/boleta", "Ver boleta por pago"),
            ("GET", "/api/v1/pagos/{pago}/boleta/descargar", "Descargar boleta por pago"),
        ],
    ),
    (
        "Servicios API",
        "Consulta pública y administración de servicios.",
        [
            ("GET", "/api/v1/servicios", "Listar servicios"),
            ("GET", "/api/v1/servicios/{servicio}", "Ver servicio"),
            ("GET", "/api/v1/admin/servicios", "Listar servicios admin"),
            ("POST", "/api/v1/admin/servicios", "Crear servicio"),
            ("GET", "/api/v1/admin/servicios/{servicio}", "Ver servicio admin"),
            ("PUT/PATCH", "/api/v1/admin/servicios/{servicio}", "Actualizar servicio"),
            ("POST", "/api/v1/admin/servicios/{servicio}/imagen", "Subir imagen"),
            ("DELETE", "/api/v1/admin/servicios/{servicio}", "Eliminar servicio"),
        ],
    ),
    (
        "Razas API",
        "Razas por especie y fotos de razas.",
        [
            ("GET", "/api/v1/razas", "Listar razas"),
            ("GET", "/api/v1/razas/{raza}", "Ver raza"),
            ("POST", "/api/v1/admin/razas", "Subir foto de raza"),
            ("DELETE", "/api/v1/admin/razas/{raza}", "Eliminar foto de raza"),
        ],
    ),
    (
        "Empleado API",
        "Operaciones internas del rol empleado.",
        [
            ("GET", "/api/v1/empleado/panel-del-dia", "Panel del día"),
            ("GET", "/api/v1/empleado/reservas", "Bandeja de reservas"),
            ("PUT/PATCH", "/api/v1/empleado/reservas/{reserva}/atender", "Atender reserva"),
            ("GET", "/api/v1/empleado/turnos", "Listar turnos"),
            ("POST", "/api/v1/empleado/turnos", "Crear turno"),
            ("GET", "/api/v1/empleado/turnos/{turno}", "Ver turno"),
            ("PUT/PATCH", "/api/v1/empleado/turnos/{turno}", "Actualizar turno"),
            ("DELETE", "/api/v1/empleado/turnos/{turno}", "Eliminar turno"),
            ("GET", "/api/v1/empleado/novedades", "Listar novedades"),
            ("POST", "/api/v1/empleado/novedades", "Crear novedad"),
            ("GET", "/api/v1/empleado/novedades/{novedad}", "Ver novedad"),
            ("PUT/PATCH", "/api/v1/empleado/novedades/{novedad}", "Actualizar novedad"),
            ("DELETE", "/api/v1/empleado/novedades/{novedad}", "Eliminar novedad"),
            ("GET", "/api/v1/empleado/ingresos", "Ver ingresos"),
        ],
    ),
    (
        "Supervisor API",
        "Ingresos, reportes y métricas.",
        [
            ("GET", "/api/v1/supervisor/ingresos", "Ver ingresos"),
            ("GET", "/api/v1/supervisor/metricas", "Ver métricas"),
            ("GET", "/api/v1/supervisor/reportes/ingresos-excel", "Descargar reporte Excel"),
        ],
    ),
    (
        "Admin API",
        "Usuarios, servicios, reservas, mascotas y configuración.",
        [
            ("GET", "/api/v1/admin/usuarios", "Listar usuarios"),
            ("POST", "/api/v1/admin/usuarios", "Crear usuario"),
            ("GET", "/api/v1/admin/usuarios/{usuario}", "Ver usuario"),
            ("PUT/PATCH", "/api/v1/admin/usuarios/{usuario}", "Actualizar usuario"),
            ("DELETE", "/api/v1/admin/usuarios/{usuario}", "Desactivar usuario"),
            ("GET", "/api/v1/admin/reservas", "Listar reservas"),
            ("GET", "/api/v1/admin/reservas/{reserva}", "Ver reserva"),
            ("PUT/PATCH", "/api/v1/admin/reservas/{reserva}", "Actualizar reserva"),
            ("DELETE", "/api/v1/admin/reservas/{reserva}", "Cancelar reserva"),
            ("GET", "/api/v1/admin/mascotas", "Listar mascotas"),
            ("POST", "/api/v1/admin/mascotas", "Crear mascota"),
            ("GET", "/api/v1/admin/mascotas/{mascota}", "Ver mascota"),
            ("PUT/PATCH", "/api/v1/admin/mascotas/{mascota}", "Actualizar mascota"),
            ("DELETE", "/api/v1/admin/mascotas/{mascota}", "Eliminar mascota"),
            ("GET", "/api/v1/admin/configuracion", "Ver configuración"),
        ],
    ),
    (
        "Otros Endpoints",
        "Servicios técnicos o globales de la API.",
        [
            ("GET", "/api/v1/health", "Verificar API activa"),
            ("GET", "/api/v1/catalogo", "Ver catálogo completo"),
            ("GET", "/api/v1/me", "Usuario actual"),
        ],
    ),
]


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_table_width(table, widths):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    for row in table.rows:
        for idx, width in enumerate(widths):
            cell = row.cells[idx]
            cell.width = Inches(width)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.first_child_found_in("w:tcW")
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(int(width * 1440)))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def set_run_font(run, size=None, bold=None, color=None, name="Calibri"):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    if size:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def add_paragraph(doc, text="", style=None):
    p = doc.add_paragraph(text, style=style)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.25
    return p


def add_endpoint_table(doc, rows):
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    table.allow_autofit = False
    headers = ["Método", "Endpoint", "Servicio"]
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        shade_cell(cell, "E8EEF5")
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER if i == 0 else WD_ALIGN_PARAGRAPH.LEFT
        r = p.add_run(h)
        set_run_font(r, size=9, bold=True, color="0B2545")
    set_repeat_table_header(table.rows[0])

    for method, endpoint, service in rows:
        cells = table.add_row().cells
        values = [method, endpoint, service]
        for i, value in enumerate(values):
            p = cells[i].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if i == 0 else WD_ALIGN_PARAGRAPH.LEFT
            r = p.add_run(value)
            set_run_font(r, size=8.5 if i == 1 else 9, bold=(i == 0), color="1F2937")
            if i == 1:
                r.font.name = "Consolas"
                r._element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")
    set_table_width(table, [0.95, 3.75, 1.65])
    doc.add_paragraph()


def configure_styles(doc):
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for style_name, size, color, before, after in [
        ("Heading 1", 16, "2E74B5", 18, 10),
        ("Heading 2", 13, "2E74B5", 14, 7),
        ("Heading 3", 12, "1F4D78", 10, 5),
    ]:
        style = styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)


def add_footer(doc):
    footer = doc.sections[0].footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = p.add_run("Pet Grooming - API REST")
    set_run_font(r, size=8.5, color="6B7280")


def build_doc():
    OUT.parent.mkdir(exist_ok=True)
    doc = Document()
    configure_styles(doc)
    add_footer(doc)

    title = doc.add_paragraph()
    title.paragraph_format.space_after = Pt(6)
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = title.add_run("Servicios API REST - Pet Grooming")
    set_run_font(run, size=22, bold=True, color="0B2545")

    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(14)
    r = subtitle.add_run("Mapa de endpoints implementados en Laravel bajo /api/v1")
    set_run_font(r, size=11, color="4B5563")

    doc.add_heading("Resumen", level=1)
    add_paragraph(
        doc,
        "Este documento lista los servicios API REST implementados en la rama codex/api-rest-metodos. "
        "La capa API está separada de las vistas Blade y usa rutas versionadas con prefijo /api/v1.",
    )

    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    headers = ["Operación", "Método HTTP", "Uso REST", "Ejemplo"]
    for idx, h in enumerate(headers):
        cell = table.rows[0].cells[idx]
        shade_cell(cell, "E8EEF5")
        p = cell.paragraphs[0]
        r = p.add_run(h)
        set_run_font(r, 9, True, "0B2545")
    set_repeat_table_header(table.rows[0])
    for row in [
        ("Consultar", "GET", "Obtener información", "/api/v1/servicios"),
        ("Registrar", "POST", "Crear o ejecutar registro", "/api/v1/reservas"),
        ("Actualizar", "PUT/PATCH", "Modificar datos", "/api/v1/clientes/mascotas/{mascota}"),
        ("Eliminar", "DELETE", "Eliminar, cancelar o desactivar", "/api/v1/empleado/turnos/{turno}"),
    ]:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            p = cells[i].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if i == 1 else WD_ALIGN_PARAGRAPH.LEFT
            r = p.add_run(value)
            set_run_font(r, 9, i == 1, "1F2937")
            if i == 3:
                r.font.name = "Consolas"
                r._element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")
    set_table_width(table, [1.2, 1.25, 2.0, 2.0])

    doc.add_heading("Diseño de endpoints", level=1)
    for text in [
        "Todas las rutas API usan el prefijo versionado /api/v1.",
        "Los recursos principales usan nombres en plural: servicios, mascotas, reservas, turnos, usuarios.",
        "Los contextos por rol están agrupados: auth, clientes, empleado, supervisor y admin.",
        "Las operaciones mantienen coherencia REST: GET consulta, POST registra, PUT/PATCH actualiza y DELETE elimina o cancela.",
    ]:
        p = doc.add_paragraph(style="List Bullet")
        r = p.add_run(text)
        set_run_font(r, 10.5, None, "1F2937")

    doc.add_heading("Listado de servicios API", level=1)
    for idx, (name, description, rows) in enumerate(GROUPS):
        if idx in (3, 6):
            doc.add_section(WD_SECTION.NEW_PAGE)
        doc.add_heading(name, level=2)
        add_paragraph(doc, description)
        add_endpoint_table(doc, rows)

    doc.add_heading("Notas de autenticación", level=1)
    for text in [
        "La API protegida usa la autenticación por sesión de Laravel, con MFA y middleware de roles.",
        "No se agregó JWT ni API Key en esta fase para no cambiar el modelo de seguridad actual del sistema.",
        "Para una aplicación móvil o frontend separado, el siguiente paso recomendado sería Laravel Sanctum.",
    ]:
        p = doc.add_paragraph(style="List Bullet")
        r = p.add_run(text)
        set_run_font(r, 10.5, None, "1F2937")

    doc.save(OUT)
    print(OUT.resolve())


if __name__ == "__main__":
    build_doc()
